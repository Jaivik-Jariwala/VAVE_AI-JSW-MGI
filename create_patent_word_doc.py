import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_charts():
    # 1. Cost Comparison Chart
    applicants = ['Student Intern', 'Large Entity (Company)']
    filing_fees = [1600, 8000]
    exam_fees = [4000, 20000]
    legal_fees = [15000, 100000] # Estimating avg legal fees

    # Totals
    totals = [sum(x) for x in zip(filing_fees, exam_fees, legal_fees)]
    
    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.bar(applicants, totals, color=['#4CAF50', '#003399'])
    
    ax.set_ylabel('Total Estimated Cost (INR)')
    ax.set_title('Patent Cost Comparison: Student vs Company')
    
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'₹{height:,}',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3),  # 3 points vertical offset
                    textcoords="offset points",
                    ha='center', va='bottom')
    
    plt.tight_layout()
    plt.savefig('cost_chart.png')
    plt.close()

    # 2. Filing Timeline Flowchart (Simplified as a Timeline Plot)
    phases = ['Provisional Filing', 'R&D / Refinement', 'Complete Specification', 'Publication', 'Examination', 'Grant']
    months = [0, 6, 12, 18, 30, 48] # Approximate timeline in months
    
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.plot(months, [1]*len(months), '-o', color='#BA0C2F', markerfacecolor='white', markersize=10, linewidth=2)
    
    for i, (phase, month) in enumerate(zip(phases, months)):
        ax.annotate(f"{phase}\n(Month {month})", 
                    xy=(month, 1), 
                    xytext=(0, 15 if i % 2 == 0 else -35), 
                    textcoords="offset points", 
                    ha='center', fontsize=9)
        
    ax.set_ylim(0.8, 1.2)
    ax.axis('off')
    ax.set_title('Indian Patent Filing Timeline', pad=20)
    plt.tight_layout()
    plt.savefig('timeline_chart.png')
    plt.close()

def create_word_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Patent Strategy & Research Report: VAVE AI System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Prepared for: JSW MG Motor India')
    doc.add_paragraph('Project: VAVE AI (Value Analysis Value Engineering) System')
    doc.add_paragraph('Purpose: Intellectual Property Protection for Internship Project')
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This report outlines the strategy for patenting the VAVE AI System. "
        "The system introduces a novel 'Physics-Informed Generative AI' approach that solves the critical problem of "
        "AI hallucination in engineering contexts. We recommend filing a 'System Patent' to cover the entire workflow, "
        "securing the invention as a corporate asset for JSW MG Motor India."
    )
    
    # 2. Scope of Protection (The Product)
    doc.add_heading('2. Scope of Protection: The "System" as a Product', level=1)
    doc.add_paragraph(
        "While computer programs 'per se' are excluded from patentability under Section 3(k) of the Indian Patent Act, "
        "a 'System' that utilizes software to achieve a technical effect is patentable. We define the VAVE AI "
        "not just as code, but as a specialized engineering system."
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Instead of Claiming...'
    hdr_cells[1].text = 'We Claim...'
    
    row = table.add_row().cells
    row[0].text = "Code for resolving car parts"
    row[1].text = "A Physics-Informed Context Resolution System comprising a data matrix and constraint enforcement module."
    
    row = table.add_row().cells
    row[0].text = "AI checking if ideas are good"
    row[1].text = "A Deterministic Autonomous Validation Engine utilizing vehicle-weight physics logic."

    # 3. Detailed Patentable Claims
    doc.add_heading('3. Detailed Patentable Inventions', level=1)
    
    doc.add_heading('A. The Physics-Informed Matrix (Context Resolver)', level=2)
    doc.add_paragraph(
        "Technical Problem: Generic LLMs do not understand engineering constraints (e.g., suggesting fuel tanks for EVs).\n"
        "Technical Solution: A pre-processing module that strictly defines the 'Ground Truth' based on a Vehicle x Component Matrix."
    )
    doc.add_paragraph("Key Code Reference: `_resolve_engineering_context`")

    doc.add_heading('B. The Autonomous Validation Engine (The Gatekeeper)', level=2)
    doc.add_paragraph(
        "Technical Problem: Generating ideas that are physically dangerous or illegal (e.g., thinning brake discs beyond safety limits).\n"
        "Technical Solution: A deterministic post-processing filter that applies hard-coded physics rules to validate AI outputs."
    )
    doc.add_paragraph("Key Code Reference: `_validate_and_filter_ideas`")

    # 4. Filing Process & Timeline
    doc.add_heading('4. Process & Timeline', level=1)
    doc.add_picture('timeline_chart.png', width=Inches(6))
    doc.add_paragraph("Figure 1: Standard Indian Patent Filing Timeline")

    # 5. Cost Analysis
    doc.add_heading('5. Cost Analysis (Estimated in INR)', level=1)
    doc.add_picture('cost_chart.png', width=Inches(5))
    doc.add_paragraph("Figure 2: Estimated Cost Breakdown (Official Fees + Legal)")
    
    doc.add_paragraph(
        "Note: As a corporate filing, JSW MG (Large Entity) fees apply. "
        "While higher than student fees, this ensures full corporate ownership and commercialization rights."
    )

    # 6. Conclusion
    doc.add_heading('6. Conclusion & Next Steps', level=1)
    doc.add_paragraph(
        "1. File Provisional Application immediately (Month 0).\n"
        "2. Refine the 'VLM Engine' and gather experimental data (Months 1-11).\n"
        "3. File Complete Specification (Month 12)."
    )

    file_name = 'VAVE_Patent_Detailed_Report.docx'
    doc.save(file_name)
    print(f"Successfully generated {file_name}")

if __name__ == "__main__":
    create_charts()
    create_word_doc()
